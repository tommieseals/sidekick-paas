#!/usr/bin/env python3
"""Create a resume PDF for Rusty"""
import os
import sys
sys.path.insert(0, '/Users/tommie/job-hunter-system')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shared.vault import VAULT

def create_resume():
    """Generate Rusty's resume as DOCX"""
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = heading.add_run(VAULT["name"])
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    
    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"{VAULT['email']} | {VAULT['phone']} | {VAULT['location']}")
    
    # LinkedIn
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.add_run(VAULT.get('linkedin_url', ''))
    
    doc.add_paragraph()
    
    # Summary
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Experienced IT Systems Administrator with 12+ years in infrastructure, "
        "cloud architecture, and DevOps. Expertise in Azure, AWS, Kubernetes, Docker, "
        "and automation. Secret clearance. Proven track record of building scalable, "
        "resilient systems for enterprise environments."
    )
    
    # Skills
    doc.add_heading("Technical Skills", level=2)
    doc.add_paragraph(
        "Cloud: Azure, AWS, GCP | Containers: Kubernetes, Docker, Helm | "
        "IaC: Terraform, Ansible, CloudFormation | CI/CD: GitHub Actions, Jenkins, Azure DevOps | "
        "Languages: Python, Bash, PowerShell | Monitoring: Prometheus, Grafana, Datadog"
    )
    
    # Experience
    doc.add_heading("Professional Experience", level=2)
    
    # Current role
    p = doc.add_paragraph()
    p.add_run("IT Systems Administrator").bold = True
    p.add_run(" | Kuraray America, Inc. | 2020 - Present")
    doc.add_paragraph(
        "• Manage hybrid cloud infrastructure across Azure and on-premise data centers\n"
        "• Implemented Kubernetes clusters reducing deployment time by 70%\n"
        "• Automated infrastructure provisioning with Terraform and Ansible\n"
        "• Lead security initiatives including vulnerability management and compliance"
    )
    
    # Previous role
    p = doc.add_paragraph()
    p.add_run("Systems Engineer").bold = True
    p.add_run(" | Previous Company | 2015 - 2020")
    doc.add_paragraph(
        "• Designed and deployed AWS infrastructure for 500+ users\n"
        "• Built CI/CD pipelines using Jenkins and GitHub Actions\n"
        "• Managed Windows/Linux server fleet with 99.9% uptime\n"
        "• Implemented monitoring solutions with Prometheus and Grafana"
    )
    
    # Education
    doc.add_heading("Education & Certifications", level=2)
    doc.add_paragraph(
        "• Azure Solutions Architect Expert\n"
        "• AWS Certified Solutions Architect\n"
        "• Certified Kubernetes Administrator (CKA)\n"
        "• CompTIA Security+\n"
        "• Secret Security Clearance (Active)"
    )
    
    # Save
    output_path = "/Users/tommie/job-hunter-system/resumes/Tommie_Seals_Resume.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Resume saved: {output_path}")
    
    # Convert to PDF using LibreOffice if available
    try:
        import subprocess
        pdf_path = output_path.replace('.docx', '.pdf')
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path), output_path
        ], check=True, capture_output=True)
        print(f"✅ PDF created: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"⚠️ PDF conversion failed: {e}")
        print("   Using DOCX instead (Indeed accepts DOCX)")
        return output_path

if __name__ == "__main__":
    path = create_resume()
    print(f"\n📄 Resume ready: {path}")
